import os
import uuid
import logging
from io import BytesIO
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status, Query, Header
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.config import settings
from app.database import get_db, SessionLocal
from models.models import Document, DocumentChunk, User
from models.schemas import DocumentResponse, DocumentFavoriteRequest
from middleware.auth_middleware import get_current_user
from services.document_processor import parse_and_chunk_document
import services.ai_service as ai_service

# Exporters
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

router = APIRouter(prefix="/documents", tags=["Documents"])
logger = logging.getLogger(__name__)

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    x_openai_key: Optional[str] = Header(None)
):
    # Validate file extension
    filename = file.filename
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    if ext not in [".pdf", ".docx", ".txt"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file format. Please upload PDF, DOCX, or TXT."
        )

    # Validate file size (20MB limit)
    contents = await file.read()
    file_size = len(contents)
    if file_size > 20 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File size exceeds maximum limit of 20MB."
        )

    # Reset file cursor after reading
    await file.seek(0)

    # Save file physically to uploads folder
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
    try:
        with open(file_path, "wb") as f:
            f.write(contents)
    except Exception as e:
        logger.error(f"Failed to save file locally: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Could not store file uploads locally."
        )

    try:
        # Extract and chunk document content
        chunked_data = parse_and_chunk_document(file_path, ext)
        if not chunked_data:
            raise ValueError("No text content could be extracted from document.")

        # Create Document record
        new_doc = Document(
            user_id=current_user.id,
            filename=filename,
            file_path=file_path,
            file_size=file_size,
            is_favorite=False
        )
        db.add(new_doc)
        await db.commit()
        await db.refresh(new_doc)

        # Generate vectors and add chunks to DB in batches
        contents = [chunk["content"] for chunk in chunked_data]
        embeddings = await ai_service.generate_embeddings_batch(contents, user_key=x_openai_key)
        
        for chunk, embedding in zip(chunked_data, embeddings):
            db_chunk = DocumentChunk(
                document_id=new_doc.id,
                content=chunk["content"],
                page=chunk["page"],
                embedding=embedding
            )
            db.add(db_chunk)
            
        await db.commit()
        
        # Format response
        return {
            "id": new_doc.id,
            "filename": new_doc.filename,
            "file_size": new_doc.file_size,
            "created_at": new_doc.created_at,
            "is_favorite": new_doc.is_favorite,
            "summary_generated": False
        }

    except Exception as e:
        logger.error(f"Error parsing uploaded document: {e}")
        # Clean up local file in case of error
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process document: {str(e)}"
        )

@router.get("", response_model=List[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document)
        .where(Document.user_id == current_user.id)
        .order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()
    
    response = []
    for doc in docs:
        response.append({
            "id": doc.id,
            "filename": doc.filename,
            "file_size": doc.file_size,
            "created_at": doc.created_at,
            "is_favorite": doc.is_favorite,
            "summary_generated": bool(doc.summary_executive)
        })
    return response

@router.get("/{doc_id}")
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return doc

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Delete physical file
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    # Delete database document (cascade will remove chunks, sessions, etc.)
    await db.delete(doc)
    await db.commit()
    return {"message": "Document deleted successfully"}

@router.put("/{doc_id}/favorite")
async def toggle_favorite(
    doc_id: int,
    req: DocumentFavoriteRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    doc.is_favorite = req.is_favorite
    db.add(doc)
    await db.commit()
    return {"message": "Favorite status updated"}

@router.post("/{doc_id}/metrics")
async def extract_document_metrics(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    x_openai_key: Optional[str] = Header(None)
):
    # Fetch doc and chunks to get sample text
    doc_result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = doc_result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunks_result = await db.execute(
        select(DocumentChunk).where(DocumentChunk.document_id == doc_id).limit(5)
    )
    chunks = chunks_result.scalars().all()
    sample_text = "\n".join([c.content for c in chunks])

    metrics = await ai_service.generate_document_metrics(sample_text, user_key=x_openai_key)
    doc.metrics = metrics
    db.add(doc)
    await db.commit()
    return metrics

@router.get("/{doc_id}/stream-summary")
async def stream_document_summary(
    doc_id: int,
    type: str = "executive",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    x_openai_key: Optional[str] = Header(None)
):
    # Get document
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Fetch chunks
    chunks_result = await db.execute(
        select(DocumentChunk).where(DocumentChunk.document_id == doc_id)
    )
    chunks = chunks_result.scalars().all()
    full_text = "\n".join([c.content for c in chunks])

    async def event_generator():
        collected_chunks = []
        async for chunk in ai_service.stream_summary(full_text, type, user_key=x_openai_key):
            collected_chunks.append(chunk)
            yield f"data: {chunk}\n\n"
        
        # Save complete summary to DB
        complete_summary = "".join(collected_chunks)
        
        # We need a new session in async context to save inside generator safely
        async with SessionLocal() as db_session:
            doc_to_update = await db_session.get(Document, doc_id)
            if doc_to_update:
                if type == "executive":
                    doc_to_update.summary_executive = complete_summary
                elif type == "detailed":
                    doc_to_update.summary_detailed = complete_summary
                else:
                    doc_to_update.summary_bullet = complete_summary
                db_session.add(doc_to_update)
                await db_session.commit()
        
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@router.get("/{doc_id}/stream-analysis")
async def stream_document_analysis(
    doc_id: int,
    category: str = "insights",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    x_openai_key: Optional[str] = Header(None)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunks_result = await db.execute(
        select(DocumentChunk).where(DocumentChunk.document_id == doc_id)
    )
    chunks = chunks_result.scalars().all()
    full_text = "\n".join([c.content for c in chunks])

    async def event_generator():
        collected_chunks = []
        async for chunk in ai_service.stream_analysis(full_text, category, user_key=x_openai_key):
            collected_chunks.append(chunk)
            yield f"data: {chunk}\n\n"
        
        complete_analysis = "".join(collected_chunks)
        async with SessionLocal() as db_session:
            doc_to_update = await db_session.get(Document, doc_id)
            if doc_to_update:
                if category == "insights":
                    doc_to_update.analysis_takeaways = complete_analysis
                elif category == "entities":
                    doc_to_update.analysis_entities = complete_analysis
                else:
                    doc_to_update.analysis_faq = complete_analysis
                db_session.add(doc_to_update)
                await db_session.commit()
        
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@router.get("/{doc_id}/stream-tool")
async def stream_ai_tool(
    doc_id: int,
    tool: str = "quiz",
    language: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    x_openai_key: Optional[str] = Header(None)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunks_result = await db.execute(
        select(DocumentChunk).where(DocumentChunk.document_id == doc_id)
    )
    chunks = chunks_result.scalars().all()
    full_pdf_text = "\n".join([c.content for c in chunks])

    async def event_generator():
        async for chunk in ai_service.stream_extra_tool(full_pdf_text, tool, language, user_key=x_openai_key):
            yield f"data: {chunk}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@router.get("/{doc_id}/search")
async def search_document(
    doc_id: int,
    q: str = Query(..., min_length=1),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # Verify doc ownership
    doc_result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = doc_result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Search for term inside chunks
    query = select(DocumentChunk).where(
        DocumentChunk.document_id == doc_id,
        DocumentChunk.content.ilike(f"%{q}%")
    ).limit(20)
    
    result = await db.execute(query)
    matching_chunks = result.scalars().all()
    
    results = []
    for chunk in matching_chunks:
        results.append({
            "text": chunk.content,
            "page": chunk.page
        })
    return results

@router.get("/{doc_id}/export")
async def export_document(
    doc_id: int,
    format: str = "pdf",
    type: str = "executive",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    summary_text = ""
    if type == "executive":
        summary_text = doc.summary_executive or "No Executive Summary generated."
    elif type == "detailed":
        summary_text = doc.summary_detailed or "No Detailed Summary generated."
    else:
        summary_text = doc.summary_bullet or "No Bullet Summary generated."

    filename = f"DocMind_Summary_{doc.filename}"

    if format == "txt":
        content = f"DocMind AI Report\n================\nDocument: {doc.filename}\nType: {type.upper()} SUMMARY\n\n{summary_text}\n"
        buffer = BytesIO(content.encode("utf-8"))
        headers = {"Content-Disposition": f"attachment; filename={filename}.txt"}
        return StreamingResponse(buffer, media_type="text/plain", headers=headers)

    elif format == "docx":
        docx_doc = docx.Document()
        docx_doc.add_heading("DocMind AI Analysis Report", 0)
        docx_doc.add_paragraph(f"Document: {doc.filename}")
        docx_doc.add_paragraph(f"Report Category: {type.upper()} SUMMARY")
        docx_doc.add_paragraph("-" * 40)
        
        # Add summary paragraphs
        lines = summary_text.split("\n")
        for line in lines:
            if line.strip().startswith("-") or line.strip().startswith("*"):
                docx_doc.add_paragraph(line.strip(), style="List Bullet")
            elif line.strip() != "":
                docx_doc.add_paragraph(line)

        buffer = BytesIO()
        docx_doc.save(buffer)
        buffer.seek(0)
        headers = {"Content-Disposition": f"attachment; filename={filename}.docx"}
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )

    elif format == "pdf":
        buffer = BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        
        styles = getSampleStyleSheet()
        # Custom stylesheet parameters
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=HexColor('#6366f1'),
            spaceAfter=20
        )
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#334155'),
            leading=14,
            spaceAfter=10
        )
        
        story = []
        story.append(Paragraph("DocMind AI Analysis Report", title_style))
        story.append(Paragraph(f"<b>Document Source:</b> {doc.filename}", body_style))
        story.append(Paragraph(f"<b>Classification Type:</b> {type.upper()} SUMMARY", body_style))
        story.append(Spacer(1, 15))
        
        # Clean formatting tags
        paragraphs = summary_text.split("\n\n")
        for para in paragraphs:
            if para.strip() != "":
                clean_para = para.replace("\n", "<br/>").replace("**", "<b>").replace("<b>", "</b>", 1) # Simple markdown cleanup
                story.append(Paragraph(clean_para, body_style))
                story.append(Spacer(1, 10))

        pdf.build(story)
        buffer.seek(0)
        headers = {"Content-Disposition": f"attachment; filename={filename}.pdf"}
        return StreamingResponse(buffer, media_type="application/pdf", headers=headers)

    else:
        raise HTTPException(status_code=400, detail="Invalid export format style.")
